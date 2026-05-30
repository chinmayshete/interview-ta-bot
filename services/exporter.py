"""
Exporter service — generates PDF and DOCX reports for interview summaries.
"""

import io
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class InterviewExporter:
    
    @staticmethod
    def generate_pdf(summary_data: dict, history: list) -> bytes:
        """Generate a professional PDF report."""
        pdf = FPDF()
        pdf.add_page()
        
        # Header
        pdf.set_font("helvetica", 'B', 16)
        pdf.cell(0, 10, "Interview Performance Report", 0, 1, 'C')
        pdf.set_font("helvetica", '', 10)
        pdf.cell(0, 10, f"Overall Score: {summary_data.get('overall_score', 'N/A')}/100", 0, 1, 'C')
        pdf.cell(0, 5, f"Rating: {summary_data.get('overall_rating', 'N/A')}", 0, 1, 'C')
        pdf.ln(10)
        
        # Summary Statement
        pdf.set_font("helvetica", 'B', 12)
        pdf.cell(0, 10, "Executive Summary", 0, 1, 'L')
        pdf.set_font("helvetica", '', 10)
        pdf.multi_cell(0, 5, summary_data.get('summary_statement', ''))
        pdf.ln(5)
        
        # Strengths & Weaknesses
        pdf.set_font("helvetica", 'B', 12)
        pdf.cell(0, 10, "Key Strengths", 0, 1, 'L')
        pdf.set_font("helvetica", '', 10)
        for s in summary_data.get('strengths', []):
            pdf.cell(0, 5, f"- {s}", 0, 1, 'L')
        pdf.ln(5)
        
        pdf.set_font("helvetica", 'B', 12)
        pdf.cell(0, 10, "Areas for Improvement", 0, 1, 'L')
        pdf.set_font("helvetica", '', 10)
        for w in summary_data.get('weaknesses', []):
            pdf.cell(0, 5, f"- {w}", 0, 1, 'L')
        pdf.ln(10)
        
        # Full Transcript
        pdf.add_page()
        pdf.set_font("helvetica", 'B', 14)
        pdf.cell(0, 10, "Detailed Interview Transcript", 0, 1, 'L')
        pdf.ln(5)
        
        for i, entry in enumerate(history):
            pdf.set_font("helvetica", 'B', 10)
            pdf.multi_cell(0, 5, f"Q{i+1}: {entry.get('question')}")
            pdf.set_font("helvetica", '', 10)
            pdf.multi_cell(0, 5, f"Answer: {entry.get('answer')}")
            pdf.set_font("helvetica", 'I', 9)
            pdf.cell(0, 5, f"Rating: {entry.get('rating')}", 0, 1, 'L')
            pdf.ln(5)
            
        return pdf.output(dest='S')

    @staticmethod
    def generate_docx(summary_data: dict, history: list) -> bytes:
        """Generate a professional DOCX report."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Interview Performance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metrics
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Overall Score: {summary_data.get('overall_score', 'N/A')}/100\n")
        run.bold = True
        p.add_run(f"Rating: {summary_data.get('overall_rating', 'N/A')}")
        
        # Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(summary_data.get('summary_statement', ''))
        
        # Strengths
        doc.add_heading('Key Strengths', level=1)
        for s in summary_data.get('strengths', []):
            doc.add_paragraph(s, style='List Bullet')
            
        # Weaknesses
        doc.add_heading('Areas for Improvement', level=1)
        for w in summary_data.get('weaknesses', []):
            doc.add_paragraph(w, style='List Bullet')
            
        # Recommendations
        doc.add_heading('Recommendation', level=1)
        doc.add_paragraph(summary_data.get('recommendation', ''))
        
        # Transcript
        doc.add_page_break()
        doc.add_heading('Detailed Interview Transcript', level=1)
        
        for i, entry in enumerate(history):
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i+1}: {entry.get('question')}")
            run.bold = True
            doc.add_paragraph(f"Answer: {entry.get('answer')}")
            doc.add_paragraph(f"Rating: {entry.get('rating')}", style='Body Text').italic = True
            
        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        return file_stream.getvalue()
